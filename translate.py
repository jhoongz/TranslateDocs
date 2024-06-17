from docx import Document
from deep_translator import GoogleTranslator

def translate_text(text, src_lang='en', dest_lang='es'):
    translator = GoogleTranslator(source=src_lang, target=dest_lang)
    return translator.translate(text)

def translate_docx(input_file, output_file, src_lang='en', dest_lang='es'):
    # Open the input Word document
    doc = Document(input_file)
    # Create a new Word document for the translated text
    translated_doc = Document()

    # Translate each paragraph
    for para in doc.paragraphs:
        if para.text.strip():  # Only translate non-empty paragraphs
            translated_text = translate_text(para.text, src_lang, dest_lang)
            translated_doc.add_paragraph(translated_text)
        else:
            translated_doc.add_paragraph('')  # Add empty paragraph if original is empty

    # Save the translated document
    translated_doc.save(output_file)
    print(f'Translation completed. Translated document saved as {output_file}')

# Example usage
input_file = 'path/to/your/input/document.docx'
output_file = 'path/to/your/output/document.docx'
translate_docx(input_file, output_file)